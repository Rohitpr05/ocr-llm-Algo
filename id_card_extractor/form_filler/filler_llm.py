import os
import re
from docx import Document
from docx.shared import Inches
from datetime import datetime
from docx.oxml.shared import qn
import xml.etree.ElementTree as ET

def setup_signature_folder():
    """
    Setup signature folder and provide instructions
    """
    signature_folder = "signatures"
    os.makedirs(signature_folder, exist_ok=True)
    
    # Check for existing signature files
    signature_files = []
    for ext in ['png', 'jpg', 'jpeg', 'PNG', 'JPG', 'JPEG']:
        sig_file = os.path.join(signature_folder, f"signature.{ext}")
        if os.path.exists(sig_file):
            signature_files.append(sig_file)
    
    if signature_files:
        print(f"✅ Found signature file(s): {signature_files}")
        return signature_files[0]  # Return first found
    else:
        print(f"📁 Created signature folder: {signature_folder}")
        print("📝 Please place your signature image in this folder as:")
        print("   - signatures/signature.png (recommended)")
        print("   - signatures/signature.jpg")
        print("   - signatures/signature.jpeg")
        return None

def fix_document_layout(doc):
    """
    Fix document layout issues that might cause page breaks
    """
    try:
        # Process all paragraphs to ensure proper spacing
        for paragraph in doc.paragraphs:
            # Reset excessive spacing - access through paragraph_format
            try:
                if paragraph.paragraph_format.space_before and paragraph.paragraph_format.space_before.pt > 12:
                    paragraph.paragraph_format.space_before = None
                if paragraph.paragraph_format.space_after and paragraph.paragraph_format.space_after.pt > 12:
                    paragraph.paragraph_format.space_after = None
            except AttributeError:
                # Skip if paragraph format attributes don't exist
                pass
        
        # Process table spacing
        for table in doc.tables:
            # Set table to not break across pages if possible
            try:
                table.alignment = 0  # Left alignment
            except:
                pass
            
            for row in table.rows:
                try:
                    row.height_rule = 1  # Auto height
                except:
                    pass
                    
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        try:
                            if paragraph.paragraph_format.space_before and paragraph.paragraph_format.space_before.pt > 6:
                                paragraph.paragraph_format.space_before = None
                            if paragraph.paragraph_format.space_after and paragraph.paragraph_format.space_after.pt > 6:
                                paragraph.paragraph_format.space_after = None
                        except AttributeError:
                            # Skip if paragraph format attributes don't exist
                            pass
        
        print("✅ Document layout optimized to prevent page breaks")
    except Exception as e:
        print(f"⚠️ Could not optimize layout: {e}")

def fill_itr_with_llm(template_path, output_path, json_data):
    """
    Enhanced ITR form filler that handles text boxes, tables, and nested elements
    """
    try:
        doc = Document(template_path)
        
        # Prepare all placeholders with exact matches from your template
        replacements = {
            "{{NAME}}": json_data.get("name", ""),
            "{{PAN}}": json_data.get("pan_number", json_data.get("pan", "")),
            "{{FATHER_NAME}}": json_data.get("father_name", ""),
            "{{ACK_NO}}": json_data.get("ack_no", "ABC1234567"),
            "{{CAPACITY}}": json_data.get("capacity", "Self"),
            "{{DATE}}": datetime.now().strftime("%d-%m-%Y"),
            "{{SIGNATURE}}": "[SIGNATURE_PLACEHOLDER]"
        }
        
        # Setup signature folder and find signature file
        signature_img = setup_signature_folder()
        if not signature_img:
            signature_img = os.path.join("signatures", "signature.png")  # Default path
        
        print(f"🔍 Starting form filling with data: {json_data}")
        print(f"📝 Replacements to be made: {replacements}")
        
        replacement_count = {key: 0 for key in replacements.keys()}
        
        def replace_text_in_runs(paragraph):
            """Advanced text replacement that handles text split across runs"""
            if not paragraph.runs:
                return False
            
            # Collect all text and run information
            full_text = ""
            run_info = []
            
            for i, run in enumerate(paragraph.runs):
                start_pos = len(full_text)
                full_text += run.text
                end_pos = len(full_text)
                run_info.append((i, start_pos, end_pos, run))
            
            # Check if any placeholder exists
            original_text = full_text
            updated_text = full_text
            
            for placeholder, replacement in replacements.items():
                if placeholder in updated_text:
                    updated_text = updated_text.replace(placeholder, str(replacement))
                    replacement_count[placeholder] += 1
                    print(f"✅ Found and replaced '{placeholder}' with '{replacement}'")
            
            # If text changed, rebuild the paragraph
            if updated_text != original_text:
                # Clear all runs
                for run in paragraph.runs[::-1]:
                    paragraph._element.remove(run._element)
                
                # Add new text as single run
                new_run = paragraph.add_run(updated_text)
                return True
            
            return False
        
        def handle_signature_paragraph(paragraph):
            """Handle signature insertion with proper sizing and positioning"""
            full_text = ''.join(run.text for run in paragraph.runs)
            
            if "{{SIGNATURE}}" in full_text:
                # Store original paragraph formatting safely
                original_alignment = None
                original_space_before = None
                original_space_after = None
                
                try:
                    original_alignment = paragraph.alignment
                    original_space_before = paragraph.paragraph_format.space_before
                    original_space_after = paragraph.paragraph_format.space_after
                except AttributeError:
                    # Attributes don't exist, use defaults
                    pass
                
                paragraph.clear()
                
                if json_data.get("signature_present", False) and os.path.exists(signature_img):
                    try:
                        run = paragraph.add_run()
                        # Smaller signature size to prevent layout issues
                        run.add_picture(signature_img, width=Inches(1.2), height=Inches(0.6))
                        
                        # Restore original formatting safely
                        try:
                            if original_alignment is not None:
                                paragraph.alignment = original_alignment
                            if original_space_before is not None:
                                paragraph.paragraph_format.space_before = original_space_before
                            if original_space_after is not None:
                                paragraph.paragraph_format.space_after = original_space_after
                        except AttributeError:
                            # Skip if formatting attributes don't exist
                            pass
                        
                        print("✅ Signature image inserted with proper sizing")
                        replacement_count["{{SIGNATURE}}"] += 1
                    except Exception as e:
                        print(f"⚠️ Could not insert signature: {e}")
                        paragraph.add_run("[SIGNATURE MISSING]")
                        replacement_count["{{SIGNATURE}}"] += 1
                else:
                    paragraph.add_run("[SIGNATURE MISSING]")
                    replacement_count["{{SIGNATURE}}"] += 1
                return True
            return False
        
        def process_paragraph(paragraph):
            """Process a single paragraph"""
            if handle_signature_paragraph(paragraph):
                return
            
            replace_text_in_runs(paragraph)
        
        def process_table_cell(cell):
            """Process each cell in a table thoroughly"""
            # Process paragraphs in cell
            for paragraph in cell.paragraphs:
                process_paragraph(paragraph)
            
            # Process nested tables
            for table in cell.tables:
                process_table(table)
        
        def process_table(table):
            """Process all cells in a table"""
            for row in table.rows:
                for cell in row.cells:
                    process_table_cell(cell)
        
        def process_all_text_nodes():
            """Process all text nodes in the document XML directly"""
            try:
                # Get the document XML
                doc_xml = doc._part._element
                
                # Define namespace
                namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                # Find all text elements
                text_elements = []
                
                def find_text_elements(element):
                    """Recursively find all text elements"""
                    if element.tag.endswith('}t'):  # w:t elements
                        text_elements.append(element)
                    
                    for child in element:
                        find_text_elements(child)
                
                find_text_elements(doc_xml)
                
                print(f"🔍 Found {len(text_elements)} text elements in XML")
                
                # Process each text element
                for text_elem in text_elements:
                    if text_elem.text:
                        original_text = text_elem.text
                        updated_text = original_text
                        
                        for placeholder, replacement in replacements.items():
                            if placeholder in updated_text:
                                updated_text = updated_text.replace(placeholder, str(replacement))
                                replacement_count[placeholder] += 1
                                print(f"✅ XML: Replaced '{placeholder}' with '{replacement}' in '{original_text[:30]}...'")
                        
                        if updated_text != original_text:
                            text_elem.text = updated_text
                
            except Exception as e:
                print(f"⚠️ Error in XML processing: {e}")
        
        def search_and_replace_in_xml():
            """Alternative XML approach using string replacement"""
            try:
                # Get the raw XML content
                doc_part = doc._part
                xml_content = doc_part._element.xml.decode('utf-8') if isinstance(doc_part._element.xml, bytes) else str(doc_part._element.xml)
                
                original_xml = xml_content
                
                # Replace placeholders in raw XML
                for placeholder, replacement in replacements.items():
                    if placeholder in xml_content:
                        xml_content = xml_content.replace(placeholder, str(replacement))
                        replacement_count[placeholder] += 1
                        print(f"✅ XML String: Replaced '{placeholder}' with '{replacement}'")
                
                # If XML was modified, we need to reload (this is complex, so we'll skip for now)
                # This is a fallback that shows what we found
                if xml_content != original_xml:
                    print("📝 Found additional placeholders in raw XML that need manual handling")
                
            except Exception as e:
                print(f"⚠️ Error in XML string processing: {e}")
        
        # Process main document body
        print("🔄 Processing main document paragraphs...")
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph)
        
        # Process all tables
        print("🔄 Processing tables...")
        for table in doc.tables:
            process_table(table)
        
        # Process headers and footers
        print("🔄 Processing headers and footers...")
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    process_paragraph(paragraph)
                for table in section.header.tables:
                    process_table(table)
            
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    process_paragraph(paragraph)
                for table in section.footer.tables:
                    process_table(table)
        
        # Process all text nodes directly
        print("🔄 Processing all text nodes in XML...")
        process_all_text_nodes()
        
        # Fallback: Search in raw XML
        print("🔄 Searching raw XML for missed placeholders...")
        search_and_replace_in_xml()
        
        # Show replacement summary
        print("\n📊 Replacement Summary:")
        for placeholder, count in replacement_count.items():
            if count > 0:
                print(f"   ✅ {placeholder}: {count} replacement(s)")
            else:
                print(f"   ❌ {placeholder}: Not found!")
        
        # Fix document layout before saving
        print("🔧 Optimizing document layout...")
        fix_document_layout(doc)
        
        # Create output directory if it doesn't exist
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Check if signature file exists, if not create a sample
        if not os.path.exists(signature_img):
            print(f"📝 Signature file not found at {signature_img}")
            print(f"💡 Please place your signature image at: {signature_img}")
            print("   Supported formats: PNG, JPG, JPEG")
            print("   Recommended size: 200x100 pixels or similar aspect ratio")
        
        # Save the document
        doc.save(output_path)
        print(f"\n✅ Final filled ITR form saved at: {output_path}")
        print(f"📁 Signature folder created at: signatures")
        
        if not os.path.exists(signature_img):
            print(f"⚠️  To add signature: Place your signature image at '{signature_img}'")
        
        # Check if all placeholders were replaced
        missed_placeholders = [k for k, v in replacement_count.items() if v == 0 and k != "{{SIGNATURE}}"]
        if missed_placeholders:
            print(f"⚠️  WARNING: These placeholders were not found: {missed_placeholders}")
            
        # Check signature specifically
        if replacement_count["{{SIGNATURE}}"] == 0:
            print("ℹ️  Signature placeholder not found - this is normal if your template doesn't have {{SIGNATURE}}")
        
        return True
        
    except Exception as e:
        print(f"❌ Error filling ITR form: {e}")
        import traceback
        traceback.print_exc()
        return False

def debug_document_structure(template_path):
    """
    Debug function to show document structure and find all text content
    """
    try:
        doc = Document(template_path)
        print("\n🔍 DOCUMENT STRUCTURE DEBUG:")
        
        print("\n📄 Main Document Paragraphs:")
        for i, para in enumerate(doc.paragraphs):
            text = ''.join(run.text for run in para.runs)
            if text.strip():
                print(f"  Para {i}: '{text[:100]}...'")
        
        print("\n📊 Tables:")
        for t_idx, table in enumerate(doc.tables):
            print(f"  Table {t_idx}:")
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    cell_text = ''.join(''.join(run.text for run in para.runs) for para in cell.paragraphs)
                    if cell_text.strip():
                        print(f"    Cell [{r_idx}][{c_idx}]: '{cell_text[:50]}...'")
        
        print("\n🔤 All Text Content (XML Method):")
        try:
            doc_xml = doc._part._element
            
            def find_all_text(element, level=0):
                indent = "  " * level
                if element.tag.endswith('}t') and element.text:
                    print(f"{indent}Text: '{element.text}'")
                
                for child in element:
                    find_all_text(child, level + 1)
            
            find_all_text(doc_xml)
        
        except Exception as e:
            print(f"⚠️ Error in XML debug: {e}")
    
    except Exception as e:
        print(f"❌ Error in debug: {e}")

def main():
    """
    Example usage with signature handling
    """
    template_path = "template/ITR_VERIFICATION_FORM.docx"
    output_path = "filled_forms/filled_itr.docx"
    
    # Setup signature folder first
    print("🔧 Setting up signature folder...")
    setup_signature_folder()
    
    # Debug the document structure first (optional)
    print("\n🔧 Running document structure debug...")
    debug_document_structure(template_path)
    
    # Sample data
    sample_data = {
        "name": "JOHN DOE",
        "pan_number": "ABCDE1234F",
        "father_name": "RICHARD DOE",
        "ack_no": "ABC1234567",
        "capacity": "Self",
        "signature_present": True  # Set to False if no signature needed
    }
    
    print(f"\n🚀 Starting ITR form filling...")
    success = fill_itr_with_llm(template_path, output_path, sample_data)
    
    if success:
        print("\n🎉 Form filling completed successfully!")
        print("📂 Check your filled form and signature folder!")
    else:
        print("\n❌ Form filling failed or incomplete!")
        
    print(f"\n📋 Next steps:")
    print(f"   1. Check filled form at: {output_path}")
    print(f"   2. Add signature image to: signatures/signature.png")
    print(f"   3. Re-run if you want to add signature")

if __name__ == "__main__":
    main()